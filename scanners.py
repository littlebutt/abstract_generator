import re
from typing import Optional

import xlrd
import docx
from docx.oxml.ns import qn
from docx.shared import Pt

import funcs
from funcs import SumFunc
from globals import Global


class ExcelScanner:
    xlsx: str = None

    def __init__(self, path: str):
        self.xlsx = xlrd.open_workbook(path)

    def read(self, row: int, col: int, sheet: Optional[str]=None) -> str:
        table = None
        if sheet is None:
            table = self.xlsx.sheet_by_name(Global.sheets[-1])
        else:
            table = self.xlsx.sheet_by_name(sheet)
        return table.cell_value(row, col)

    def get_sheets(self):
        return self.xlsx.sheet_names()


class WordScanner:

    def __init__(self, path: str, dst: str):
        self.docx = docx.Document(path)
        self.name = dst

    def call_func(self, matched):
        target = matched.group(1)
        target = target[1:-1]
        if 'SUM' in target:
            target = target.strip()
            target = target[3:]
            target = target[1:-1]
            row, col = target.split(',')
            sum = SumFunc(Global.excel_scanner)
            return sum.run(row=row, col=col, sheets=Global.sheets)
        else:
            target = target.strip()
            target_params = target.split(',')
            if len(target_params) == 2:
                row, col = target.split(',')
                norm = funcs.NormFunc(Global.excel_scanner)
                return norm.run(row=row, col=col)
            elif len(target_params) == 3:
                row, col, sheet = target.split(',')
                norm = funcs.NormFunc(Global.excel_scanner)
                return norm.run(row=row, col=col, sheet=sheet)

    def scan(self):
        for paragraph in self.docx.paragraphs:
            while re.search('\{.+}', paragraph.text):
                res = re.sub('(\{.+?})', self.call_func, paragraph.text)
                paragraph.text = res
                for run in paragraph.runs:
                    run.font.name = Global.paragraph_font_family
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), Global.paragraph_font_family)
                    run.font.size = Pt(Global.paragraph_font_size)
        for table in self.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    while re.search('\{.+}', cell.text):
                        res = re.sub('(\{.+?})', self.call_func, cell.text)
                        cell.text = res
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = Global.table_font_family
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), Global.table_font_family)
                                run.font.size = Pt(Global.table_font_size)
        self.docx.save(self.name)
