import re

import docx
from docx.oxml.ns import qn
from docx.shared import Pt

import funcs
from funcs import SumFunc
from globals import Global


class WordWriter:

    def __init__(self, path: str, dst: str):
        self.docx = docx.Document(path)
        self.name = dst
        self.func_list = [SumFunc, funcs.AvgFunc, funcs.NormFunc]

    def call_func(self, matched):
        target = matched.group(1)
        target = target[1:-1]
        for func in self.func_list:
            if func.match(target):
                f = func(Global.excel_scanner)
                return f.run(target)

    def scan(self):
        for paragraph in self.docx.paragraphs:
            while re.search('{.+}', paragraph.text):
                res = re.sub('({.+?})', self.call_func, paragraph.text)
                paragraph.text = res
                for run in paragraph.runs:
                    run.font.name = Global.paragraph_font_family
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), Global.paragraph_font_family)
                    run.font.size = Pt(Global.paragraph_font_size)
        for table in self.docx.tables:
            for row in table.rows:
                for cell in row.cells:
                    while re.search('{.+}', cell.text):
                        res = re.sub('({.+?})', self.call_func, cell.text)
                        cell.text = res
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = Global.table_font_family
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), Global.table_font_family)
                                run.font.size = Pt(Global.table_font_size)
        self.docx.save(self.name)